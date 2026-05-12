#!/usr/bin/env python3
"""Genera B370_Estrategia_Cliente.docx para compartir con el equipo."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'data', 'B370_Estrategia_Cliente.docx')

# ── Colores ──
VERDE   = RGBColor(0x00, 0x6A, 0x4E)
DORADO  = RGBColor(0xFF, 0xD7, 0x00)
OSCURO  = RGBColor(0x0D, 0x1B, 0x2A)
GRIS    = RGBColor(0x55, 0x55, 0x55)
ROJO    = RGBColor(0xC0, 0x39, 0x2B)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)
VERDE_L = RGBColor(0xE8, 0xF5, 0xE9)
GRIS_L  = RGBColor(0xF5, 0xF5, 0xF5)

doc = Document()

# ── Márgenes ──
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──
def add_heading(text, level=1, color=VERDE, size=18, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_subheading(text, color=OSCURO, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_body(text, size=10.5, color=GRIS, space_after=4, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p

def add_bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRIS

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '006A4E')
    pBdr.append(bottom)
    pPr.append(pBdr)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, header_color='006A4E', col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLANCO
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = OSCURO
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run('B370 LÍNEA DEPORTIVA')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = VERDE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Documento Estratégico de Marca y Cliente')
r2.font.size = Pt(14)
r2.font.color.rgb = GRIS

fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = fecha.add_run('Mayo 2026  |  Confidencial — Solo equipo interno')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r3.italic = True

doc.add_paragraph()
add_divider()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.paragraph_format.space_before = Pt(12)
r4 = intro.add_run(
    'Este documento contiene los perfiles de cliente, estrategia de contenido, '
    'voz de marca, KPIs y secuencias de email de B370. '
    'Es la guía operativa de todo el equipo.'
)
r4.font.size = Pt(10.5)
r4.font.color.rgb = GRIS
r4.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════
#  SECCIÓN 1 — BUYER PERSONAS
# ══════════════════════════════════════════════
add_heading('1. Perfiles de Cliente (Buyer Personas)', level=1, size=20)
add_body(
    'Identificados a través de la metodología "7 Maletas" de Felipe Vergara. '
    'Toda pieza de comunicación debe apuntar a una de estas personas.',
    size=10.5
)
add_divider()

# ── Persona 1 ──
add_subheading('Persona 1 — Premium Street/Stadium Buyer', color=VERDE, size=14)
add_body('Hombre, 20-35 años | Poder adquisitivo medio-alto | Hincha activo', bold=True, size=10)

add_subheading('¿Quién es?', size=11)
add_body(
    'Urbano, trabajador, ingresos propios. Va al estadio o sigue los partidos en bares. '
    'La camiseta es identidad y look — no solo apoyo al equipo. Le importa verse real.',
    size=10.5
)

add_subheading('¿Qué compra en B370?', size=11)
add_bullet('Tipo Original ($109.900) o 1.1 ($119.900)')
add_bullet('Selecciones: Colombia, Argentina, Brasil')
add_bullet('Clubes top: Atlético Nacional, Barcelona, Real Madrid')

add_subheading('Motivadores principales', size=11)
add_table(
    ['#', 'Motivador', 'Qué activa'],
    [
        ['1', 'Estatus social', 'Verse bien, tener "la buena"'],
        ['2', 'Lealtad al equipo', 'Vestir bien al equipo del corazón'],
        ['3', 'Pertenencia', 'Sentirse parte de la tribu hincha'],
    ],
    col_widths=[1.2, 4, 9]
)

add_subheading('Objeciones típicas', size=11)
add_bullet('¿Es realmente 1.1 o es Tipo Fan disfrazada? → Resolver con video del producto')
add_bullet('¿Le ponen los parches oficiales? → Foto del detalle en zoom')
add_bullet('¿La talla es real? → Tabla + testimonios con medidas')
add_bullet('¿No será mejor pedir la oficial directo del equipo? → Comparativo de precio')

add_subheading('Cómo hablarle', size=11)
add_bullet('Tono: energético, con jerga futbolera natural')
add_bullet('Referencias a estadio, partido, clásico')
add_bullet('Mostrar el detalle visual: parches, sublimación, hilo')
add_bullet('Apelar al orgullo: "Esta es la que se va al estadio"')
add_bullet('Contenido de credibilidad funciona muy bien (Jorman Campuzano, sede física)')

add_subheading('Canales principales', size=11)
add_bullet('Instagram (principal) — Reels de detalle de producto')
add_bullet('TikTok — contenido aspiracional, partidos, looks')
add_bullet('Meta Ads: intereses equipo + edad 20-35 + género masculino')

add_divider()

# ── Persona 2 ──
add_subheading('Persona 2 — Budget Fan Buyer', color=VERDE, size=14)
add_body('Hincha promedio, 18-50 años | Presupuesto ajustado | Todo Colombia', bold=True, size=10)

add_subheading('¿Quién es?', size=11)
add_body(
    'Hincha real, quiere la camiseta del equipo sin pagar precio oficial. '
    'Compra desde su ciudad o municipio. Valora mucho poder pagar al recibir '
    'porque no tiene tarjeta o no confía en pago online.',
    size=10.5
)

add_subheading('¿Qué compra en B370?', size=11)
add_bullet('Tipo Fan ($79.900) o Retro ($79.900)')
add_bullet('A veces compra para regalar: hijo, papá, amigo')
add_bullet('Equipos colombianos: Nacional, Millonarios, Junior, América')
add_bullet('Selección Colombia (principal motivo de compra en 2026)')

add_subheading('Motivadores principales', size=11)
add_table(
    ['#', 'Motivador', 'Qué activa'],
    [
        ['1', 'Pertenencia', 'Sentirse hincha legítimo sin gastar de más'],
        ['2', 'Valor por precio', 'Sacarle el jugo al billete'],
        ['3', 'Honor familiar', 'Regalar al hijo, al papá, al sobrino'],
    ],
    col_widths=[1.2, 4, 9]
)

add_subheading('Objeciones típicas', size=11)
add_bullet('¿No será una falsificación que se cae a pedazos? → Sede física + WhatsApp humano')
add_bullet('¿Cobran envío? → Copy claro del umbral de envío gratis ($200.000)')
add_bullet('¿De verdad puedo pagar al recibir? → Video explicativo del proceso')
add_bullet('¿Llega a mi municipio? → Lista de cobertura a nivel nacional')

add_subheading('Cómo hablarle', size=11)
add_bullet('Tono: cálido, cercano, sin tecnicismos')
add_bullet('Apelar al sentimiento de hincha y al regalo familiar')
add_bullet('Mostrar la camiseta en uso real (no fotos editoriales frías)')
add_bullet('Reforzar contra-entrega y WhatsApp como red de seguridad emocional')
add_bullet('Lenguaje sencillo: precio claro, talla clara, plazo claro')

add_subheading('Canales principales', size=11)
add_bullet('Facebook (más que Instagram) — sigue páginas del equipo')
add_bullet('WhatsApp — canal de servicio donde más cómodo se siente')
add_bullet('Meta Ads: equipo + ubicación geográfica + comportamiento de compra online')

add_divider()

# ── Personas auxiliares ──
add_subheading('Personas auxiliares (tratamiento secundario)', color=GRIS, size=13)

add_subheading('Padres de niños en escuelas de fútbol', size=11)
add_body('Padre/madre, 30-50 años. Compra para el hijo. Motivador: orgullo del hijo deportista + durabilidad + precio razonable.', size=10.5)
add_body('Hablarle: tono seguro y empático, apelar al orgullo del hijo y la foto del partido, resaltar calidad de lavado.', size=10.5)

add_subheading('Comprador de regalo', size=11)
add_body('Amigo/familiar/pareja del hincha. Sabe poco de camisetas. Motivador: acertar en el regalo.', size=10.5)
add_body('Hablarle: tono guía y didáctico, ofrecerle ayuda activa ("dime el equipo y te recomiendo"), resaltar bono regalo.', size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════
#  SECCIÓN 2 — VOZ Y TONO
# ══════════════════════════════════════════════
add_heading('2. Voz y Tono de B370', level=1, size=20)
add_body('Si la pieza no suena así, no es B370. Regla de oro: léelo en voz alta. Si no lo dirías hablando con un parcero hincha, no va.', size=10.5)
add_divider()

add_subheading('Atributos de voz', size=12)
add_table(
    ['Atributo', 'SÍ es B370', 'NO es B370'],
    [
        ['Pasional', 'Habla con emoción real del fútbol', 'Frío, descriptivo, "informativo"'],
        ['Auténtico', 'Slang colombiano natural ("crack", "parcero", "berraco")', 'Inventado, forzado, traducido'],
        ['Cercano', 'Habla de "tú", como amigo', '"Estimado cliente", "ustedes"'],
        ['Aspiracional', 'Conecta camiseta con identidad y orgullo', 'Solo ficha técnica'],
        ['Comercialmente directo', 'CTA claro, sin rodeos en momentos de venta', 'Pasivo, "si gustas", "tal vez"'],
    ],
    col_widths=[4, 6.5, 6.5]
)

add_subheading('Reglas de redacción', size=12)
add_bullet('Habla de "tú" siempre. Nunca "usted" en orgánico ni en pauta.')
add_bullet('Frases cortas. Máximo 15 palabras. Mejor dos frases que una larga.')
add_bullet('CTAs en imperativo amable: "Pídela ya", "Mírala aquí", "Llévatela".')
add_bullet('Referencias futboleras cuando aplique: "para el partido del domingo", "para llevarla al estadio".')
add_bullet('Cierre comercial obligatorio: "En B370, vestimos la pasión." (solo en piezas de venta directa).')
add_bullet('Emojis con criterio: ⚽🔥🇨🇴🥇 sí. Emojis de llanto o corazón roto: no.')

add_subheading('Tono según contexto', size=12)
add_table(
    ['Contexto', 'Energía', 'Ejemplo'],
    [
        ['Lanzamiento / pre-venta',
         'Alta | Urgente | Directo',
         '"Mañana 10am. La nueva Selección Colombia 2026. Solo 50 unidades con cupón VIPB370. Pon la alarma."'],
        ['Educación / orgánico',
         'Media | Explicativo',
         '"¿Sabes la diferencia entre Tipo Fan y Tipo Original? Te la explico en 30 segundos para que escojas la tuya."'],
        ['Servicio / WhatsApp',
         'Cálida | Humana | Resolutiva',
         '"¡Hola crack! Claro que sí, te confirmo: la talla M de la Nacional ya está, te la enviamos hoy y la pagas al recibir. ¿La marco?"'],
        ['Comentario negativo',
         'Respetuosa | Sin defensividad',
         '"Hola, lamento mucho lo que pasó. Te escribo por DM ahora mismo para resolverlo bien. Tu compra importa."'],
        ['Aspiracional / storytelling',
         'Emotiva | Orgullosa',
         '"Esta camiseta no se queda en el clóset. Se va al partido, al asado, al colegio del hijo. Es identidad."'],
    ],
    col_widths=[4.5, 3.5, 9]
)

add_subheading('Vocabulario', size=12)
add_body('SÍ decimos:', bold=True, size=10.5)
add_bullet('"Crack", "parcero", "berraco" (con moderación)')
add_bullet('"Camiseta" (no "jersey" en orgánico)')
add_bullet('"Tipo Fan", "Tipo Original", "1.1", "Retro" (nombres oficiales de calidades)')
add_bullet('"Pagas al recibir", "contra-entrega"')
add_bullet('"Hincha", "fanático", "coleccionista" según audiencia')

add_body('NO decimos:', bold=True, size=10.5)
add_bullet('"Producto" → decimos "camiseta", "buzo", "gabán"')
add_bullet('"Cliente" en orgánico → decimos "tú", "vos", o el nombre')
add_bullet('"Inventario", "stock disponible" → decimos "ya está", "te queda", "última talla"')
add_bullet('"Shop now", "limited edition" (inglés innecesario para Colombia)')

doc.add_page_break()

# ══════════════════════════════════════════════
#  SECCIÓN 3 — ESTRATEGIA DE CONTENIDO
# ══════════════════════════════════════════════
add_heading('3. Estrategia de Contenido', level=1, size=20)
add_body('Post-launch comercial agresivo. El foco: convertir tráfico en pedidos, generar prueba social y manejar objeciones.', size=10.5)
add_divider()

add_subheading('Pilares de contenido', size=12)
add_table(
    ['Pilar', '%', 'Para qué sirve', 'Ejemplos'],
    [
        ['Venta directa', '35%', 'Conversión inmediata', 'Producto + precio + CTA a WA o web'],
        ['Educación', '25%', 'Resolver objeciones', '¿Cómo elegir talla? ¿Qué diferencia calidades?'],
        ['Entretenimiento', '20%', 'Viralidad y alcance', 'Reacciones a partidos, trends fútbol colombiano'],
        ['Aspiracional', '10%', 'Identidad de marca', 'Hincha en contexto real, orgullo colombiano'],
        ['Credibilidad', '10%', 'Prueba social y confianza', 'Jorman Campuzano, sede La Ceja, testimonios'],
    ],
    col_widths=[4, 1.5, 4.5, 7]
)

add_subheading('Canales y su rol', size=12)
add_table(
    ['Canal', 'Pilar dominante', 'Frecuencia', 'KPI principal'],
    [
        ['Instagram (feed + reels)', 'Venta + Aspiracional', '4-5 publicaciones/semana', 'Alcance + clics al perfil'],
        ['Instagram Stories', 'Venta + Servicio', 'Diario', 'Respuestas + clics al sticker'],
        ['TikTok', 'Entretenimiento + Educación', '3-4/semana', 'Reproducciones + saves'],
        ['Facebook', 'Venta + Credibilidad', '2-3/semana', 'Mensajes a Messenger'],
        ['WhatsApp Status', 'Venta directa + cupones', 'Diario', 'Mensajes entrantes'],
        ['WhatsApp Broadcast', 'Conversión', 'Semanal', 'Tasa de respuesta'],
        ['Email (Omnisend)', 'Carrito abandonado, post-compra', 'Automatizado', 'Apertura + revenue por email'],
    ],
    col_widths=[5, 4.5, 4, 4]
)

add_subheading('Las 5 fuentes de ideas', size=12)
add_body('Toda idea de contenido sale de una de estas fuentes. Si no, sospecha.', size=10.5)
rows_fuentes = [
    ['1', 'WhatsApp de B370', 'LA MINA DE ORO. Cada pregunta repetida 3+ veces = pieza de contenido'],
    ['2', 'Comentarios y DMs', 'Objeciones públicas → piezas que las resuelven públicamente'],
    ['3', 'Banco de credibilidad', 'Cada hito de B370: lanzamiento, Jorman, cobertura nacional, etc.'],
    ['4', 'Calendario deportivo', 'El fútbol es predecible. Marcar 30 días antes: finales, clásicos, eliminatorias, Mundial'],
    ['5', 'Errores y aprendizajes', 'Convertir el "se nos cayó algo" en historia humana de marca'],
]
add_table(['#', 'Fuente', 'Cómo usarla'], rows_fuentes, col_widths=[1.2, 4.5, 11.3])

add_subheading('Reglas del contenido', size=12)
add_bullet('Si no aporta valor o no vende, no se publica. Cero "fillers".')
add_bullet('Cada pieza tiene un solo objetivo: vender, educar o entretener. No las tres a la vez.')
add_bullet('CTA explícito siempre que sea pieza de venta. WhatsApp o web, nunca ambiguo.')
add_bullet('No publicar el día siguiente de una tragedia deportiva sin sensibilidad editorial.')
add_bullet('No saltarse el calendario de Beto — si hay un lanzamiento agendado, todo se alinea.')

add_subheading('Estrategia de reciclaje (una idea = 3+ plataformas)', size=12)
add_bullet('Reel de Instagram → TikTok (sin watermark) + Facebook Reels')
add_bullet('Carrusel de IG → 3 stories sucesivas + post de Facebook')
add_bullet('Foto de producto → feed IG + status WhatsApp + email + creatividad Meta Ads')
add_bullet('Pregunta frecuente del WA → carrusel + reel + email + post blog (SEO)')

doc.add_page_break()

# ══════════════════════════════════════════════
#  SECCIÓN 4 — SECUENCIAS DE EMAIL
# ══════════════════════════════════════════════
add_heading('4. Secuencias de Email Automáticas (Omnisend)', level=1, size=20)
add_body('Estas 3 secuencias deben activarse en Omnisend. Son las de mayor impacto en conversión.', size=10.5)
add_divider()

# ── Secuencia 1: Carrito abandonado ──
add_subheading('Secuencia 1 — Carrito Abandonado', color=ROJO, size=14)
add_body('Trigger: cliente agrega producto al carrito pero no completa la compra. Espera 1h antes del primer email.', size=10.5)

add_table(
    ['Email', 'Timing', 'Asunto', 'Preheader', 'CTA'],
    [
        ['#1', '1h después',
         '¿Se te quedó algo, crack? 🔥',
         'Tu camiseta te está esperando...',
         'Completar mi pedido'],
        ['#2', '24h después',
         'Última vez que te lo digo 👀',
         'No queremos que te quedes sin ella',
         'La quiero ahora'],
        ['#3', '72h después',
         'Stock bajo en tu talla — avísate',
         'Quedan pocas. En serio.',
         'Asegurar mi talla'],
    ],
    col_widths=[1.5, 2.5, 5, 5, 3.5]
)

add_subheading('Copy Email #1 — 1 hora después', size=11)
add_body('Asunto: ¿Se te quedó algo, crack? 🔥', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Vimos que estuviste mirando [producto] y no terminaste el pedido.\n\n'
    'Tranquilo, pasa. Pero te decimos una cosa: esa camiseta ya la tiene alguien en el carrito.\n\n'
    'Si la quieres, pídela ya. Llega en 2-5 días y la pagas al recibir — sin tarjeta, sin riesgo.\n\n'
    '¿Dudas? Escríbenos por WhatsApp ahora mismo: wa.me/573218159715\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_subheading('Copy Email #2 — 24 horas después', size=11)
add_body('Asunto: Última vez que te lo digo 👀', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Ayer te dejaste [producto] en el carrito. Hoy te la volvemos a mostrar.\n\n'
    'No te pedimos que te apures. Solo te pedimos que no te vayas a arrepentir cuando la veas agotada.\n\n'
    'Paga al recibir. Llega a todo Colombia. Sin excusas.\n\n'
    '¿Tienes una duda de talla? Mándanos mensaje: estamos en WhatsApp.\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_subheading('Copy Email #3 — 72 horas después', size=11)
add_body('Asunto: Stock bajo en tu talla — avísate', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Revisamos el stock de [producto] y quedan pocas unidades en las tallas más pedidas.\n\n'
    'No inventamos urgencia — te avisamos cuando es real. Y esta vez es real.\n\n'
    'Si la quieres, el momento es ahora. Después no respondemos por las tallas.\n\n'
    'Paga al recibir. Envío a todo Colombia. Atención humana por WhatsApp.\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_divider()

# ── Secuencia 2: Post-compra ──
add_subheading('Secuencia 2 — Post-compra / Bienvenida al Cliente', color=VERDE, size=14)
add_body('Trigger: pedido confirmado en WooCommerce. Objetivo: generar confianza, reducir ansiedad post-compra, conseguir review.', size=10.5)

add_table(
    ['Email', 'Timing', 'Asunto', 'Preheader', 'CTA'],
    [
        ['#1', 'Inmediato',
         '¡Pedido recibido! Ya lo estamos alistando 🔥',
         'Tu camiseta empieza el viaje hoy',
         'Ver mi pedido'],
        ['#2', 'Día de despacho',
         'Tu camiseta salió de La Ceja ⚽',
         'Ya va en camino, crack',
         'Rastrear mi envío'],
        ['#3', 'D+7 (post-entrega)',
         '¿Cómo te quedó? Cuéntanos 🏆',
         '30 segundos y nos ayudas un montón',
         'Dejar mi reseña'],
        ['#4', 'D+14',
         'La próxima ya la tienes elegida?',
         'Mira lo que llegó nuevo esta semana',
         'Ver novedades'],
    ],
    col_widths=[1.5, 2.5, 5.5, 5, 3]
)

add_subheading('Copy Email #1 — Confirmación inmediata', size=11)
add_body('Asunto: ¡Pedido recibido! Ya lo estamos alistando 🔥', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    '¡Buena elección, crack! Tu pedido #[número] ya está en nuestras manos.\n\n'
    'Lo que pediste: [producto] — Talla [talla]\n'
    'Lo que pasa ahora: lo alistamos y lo despachamos desde La Ceja, Antioquia.\n'
    'Lo que recibes: tu camiseta en 2-5 días hábiles. Pagas al recibir.\n\n'
    '¿Alguna duda? Estamos en WhatsApp: wa.me/573218159715\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_subheading('Copy Email #3 — Solicitud de review (D+7)', size=11)
add_body('Asunto: ¿Cómo te quedó? Cuéntanos 🏆', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Ya deberías tener tu [producto] en las manos. Esperamos que te haya quedado perfecta.\n\n'
    'Un favor: ¿nos dejas tu opinión en Google? 30 segundos. No más.\n\n'
    'Tu reseña le ayuda a otros hinchas a animarse a pedir la suya — y nos ayuda a nosotros a seguir mejorando.\n\n'
    '[Botón: Dejar mi reseña]\n\n'
    'Si algo no estuvo bien, escríbenos ANTES de publicar — lo resolvemos, en serio.\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_divider()

# ── Secuencia 3: Reactivación ──
add_subheading('Secuencia 3 — Reactivación de Clientes Inactivos (60+ días)', color=OSCURO, size=14)
add_body('Trigger: cliente con compra previa que no ha vuelto en 60+ días. Objetivo: reactiva o limpia la lista.', size=10.5)
add_body('Segmento: "Clientes — última compra hace 60+ días" en Omnisend.', size=10, bold=True)

add_table(
    ['Email', 'Timing', 'Asunto', 'Preheader', 'CTA'],
    [
        ['#1', 'Día 0',
         'Te extrañamos, [nombre] ⚽',
         'Han pasado 2 meses. ¿Todo bien?',
         'Ver lo nuevo'],
        ['#2', 'Día 5',
         'Mira lo que llegó desde que te fuiste',
         'Colombia Mundial, Nacional temporada nueva...',
         'Explorar novedades'],
        ['#3', 'Día 12',
         'Última oportunidad: 10% OFF para ti',
         'Solo hasta el domingo. Sin excusas.',
         'Usar mi descuento'],
        ['#4', 'Día 18',
         'Antes de decir adiós...',
         '¿Seguimos en contacto o te sacamos de la lista?',
         'Sí, sigo aquí'],
    ],
    col_widths=[1.5, 2.5, 5.5, 5, 3]
)

add_subheading('Copy Email #1 — Reactivación suave', size=11)
add_body('Asunto: Te extrañamos, [nombre] ⚽', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Hace un tiempo que no nos pedías nada. Esperamos que estés bien — y que tu camiseta de [equipo] siga dando guerra.\n\n'
    'Desde tu última compra llegaron cosas nuevas: la Selección Colombia para el Mundial 2026, '
    'la temporada nueva de Atlético Nacional, retros que no habíamos tenido nunca.\n\n'
    'Échales un ojo. Si ves algo que te gusta, ya sabes: pagas al recibir, llega a todo Colombia.\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_subheading('Copy Email #3 — Cupón de reactivación', size=11)
add_body('Asunto: Última oportunidad: 10% OFF para ti', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'Te tenemos algo: 10% de descuento en tu próximo pedido.\n\n'
    'Código: VUELVE10\n'
    'Válido hasta el domingo.\n\n'
    'Sin mínimo de compra. Sin letra chiquita. Solo para ti porque ya eres parte de B370.\n\n'
    '¿Qué esperas? Esa camiseta que te estabas mirando ya tiene descuento.\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

add_subheading('Copy Email #4 — Último intento (win-back o baja)', size=11)
add_body('Asunto: Antes de decir adiós...', bold=True, size=10)
add_body(
    'Hola [nombre],\n\n'
    'No te vamos a mandar más correos si no quieres. Serios.\n\n'
    'Solo queremos saber: ¿seguimos en contacto o prefieres que te saquemos de la lista?\n\n'
    'Si te quedas, te enterarás primero de lanzamientos, descuentos y novedades del catálogo.\n\n'
    'Si no, también está bien. Sin dramas.\n\n'
    '[Botón principal: Sí, sigo aquí]\n'
    '[Link texto: Prefiero que me den de baja]\n\n'
    'En B370, vestimos la pasión.',
    size=10.5, color=OSCURO
)

doc.add_page_break()

# ══════════════════════════════════════════════
#  SECCIÓN 5 — KPIs
# ══════════════════════════════════════════════
add_heading('5. KPIs y Métricas Clave', level=1, size=20)
add_body('Sin métricas claras, cualquier campaña parece exitosa. Estos son los números que importan.', size=10.5)
add_divider()

add_subheading('KPIs comerciales (los que le importan a Beto)', size=12)
add_table(
    ['KPI', 'Objetivo Q2 2026', 'Dónde se ve'],
    [
        ['Pedidos / mes', 'Definir baseline primeros 30 días', 'WooCommerce + Quenti'],
        ['Ticket promedio', '$130.000+ COP', 'WooCommerce'],
        ['Tasa de conversión web', '1.5%+', 'Google Analytics 4'],
        ['Carrito abandonado recuperado', '15%+', 'Omnisend'],
        ['Costo de adquisición (CAC)', '< $25.000 COP', 'Meta Ads Manager'],
        ['ROAS Meta Ads', '> 3.0x', 'Meta Ads Manager'],
    ],
    col_widths=[6, 5, 6]
)

add_subheading('KPIs de contenido orgánico', size=12)
add_table(
    ['KPI', 'Plataforma', 'Objetivo'],
    [
        ['Alcance mensual', 'Instagram', 'Crecimiento +15% mes a mes'],
        ['Alcance mensual', 'TikTok', 'Crecimiento +25% mes a mes'],
        ['Tasa de guardados', 'Instagram Reels', '> 3%'],
        ['Mensajes entrantes', 'WhatsApp', 'Crecimiento +20% mes a mes'],
        ['Engagement rate', 'IG + TikTok', '> 5%'],
        ['Crecimiento seguidores', 'IG + TikTok (combinado)', '+500/mes'],
    ],
    col_widths=[5, 4, 8]
)

add_subheading('KPIs de servicio', size=12)
add_table(
    ['KPI', 'Objetivo'],
    [
        ['Tiempo de respuesta WhatsApp (horario laboral)', '< 30 minutos'],
        ['Tiempo de respuesta WhatsApp (fuera de horario)', '< 12 horas'],
        ['Tasa de cierre de venta por WhatsApp', '> 30%'],
        ['Reviews Google nuevas / mes', '10+ (superar vulnerabilidad #1: cero reseñas)'],
        ['NPS post-compra', '> 50'],
    ],
    col_widths=[10, 7]
)

add_subheading('Ritmo de revisión', size=12)
add_table(
    ['Frecuencia', 'Qué se revisa', 'Responsable'],
    [
        ['Diario', 'Pedidos del día + mensajes WhatsApp pendientes', 'Equipo operativo'],
        ['Semanal', 'Performance de pauta + contenido top de la semana', 'Juanjo'],
        ['Quincenal', 'Ajuste de creatividades y audiencias en Meta Ads', 'Juanjo'],
        ['Mensual', 'Reporte completo a Beto + plan siguiente mes', 'Juanjo'],
        ['Trimestral', 'Revisión estratégica, ajuste de objetivos', 'Juanjo + Beto'],
    ],
    col_widths=[3.5, 9.5, 4]
)

add_subheading('Reglas para reportar a Beto', size=12)
add_bullet('Siempre comparar contra mes anterior, no contra ideales aspiracionales.')
add_bullet('Mostrar pesos antes que porcentajes — "ingresamos $X con $Y de inversión" pesa más que "ROAS 3.2x".')
add_bullet('Una página máximo por reporte mensual. Si no cabe, no se va a leer.')
add_bullet('Cada reporte termina con la decisión que se tomará el mes siguiente — no datos sin acción.')
add_bullet('Si algo va mal, se dice claro y se propone la corrección. No se esconde.')

# ── Footer ──
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(40)
r_f = footer_p.add_run('B370 Línea Deportiva — La Ceja, Antioquia — b370sports.com')
r_f.font.size = Pt(10)
r_f.font.color.rgb = GRIS
r_f.italic = True

footer_p2 = doc.add_paragraph()
footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f2 = footer_p2.add_run('WhatsApp atención: +57 321 815 9715  |  Documento confidencial — solo equipo interno')
r_f2.font.size = Pt(9)
r_f2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r_f2.italic = True

# ── Guardar ──
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f'Generado: {os.path.abspath(OUTPUT)}')
