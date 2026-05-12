#!/usr/bin/env python3
"""Genera B370_Informe_Ventas_Mayo2026.docx con datos reales de Quenti + NEXY."""
import pandas as pd
import warnings
warnings.filterwarnings('ignore')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(__file__)
OUTPUT = os.path.join(BASE, '..', 'data', 'B370_Informe_Ventas_Mayo2026.docx')

# ══ Cargar datos ══════════════════════════════════════════════════════
df = pd.read_excel(os.path.join(BASE, '..', 'data', 'clientes',
    'HISTORIAL DE COMPRAS X PRODUCOT Y CLIENTE CUENTI DESDE 040426 HASTA 060526.xlsx'))
df['nombre'] = df['# Transaccion'].astype(str).str.strip()
df['cantidad'] = pd.to_numeric(df['Impoconsumo + ICUI'], errors='coerce').fillna(0)
df['total'] = pd.to_numeric(df['Total'], errors='coerce').fillna(0)
df['fecha'] = pd.to_datetime(df['Fecha Registro.1'], errors='coerce')
df['fecha_dia'] = df['fecha'].dt.date
df['dia_semana'] = df['fecha'].dt.day_name()
df['hora'] = df['fecha'].dt.hour

DIAS_ES = {'Monday':'Lunes','Tuesday':'Martes','Wednesday':'Miércoles',
           'Thursday':'Jueves','Friday':'Viernes','Saturday':'Sábado','Sunday':'Domingo'}

cam_mask = df['nombre'].str.upper().str.contains('CAMISETA|BUZO|GABAN|POLO|ENTRENO', na=False)
df_cam = df[cam_mask].copy()

def get_calidad(n):
    n = n.upper()
    if '1,1' in n or '1.1' in n: return '1.1'
    if 'TIPO ORIGINAL' in n: return 'Tipo Original'
    if 'FAN' in n: return 'Tipo Fan'
    if 'RETRO' in n: return 'Retro'
    return 'Otra'

def get_equipo(n):
    n = n.upper()
    if 'COLOMBIA' in n: return 'Colombia'
    if 'NACIONAL' in n or 'ATLETICO' in n: return 'Atlético Nacional'
    if 'MILLONARIOS' in n: return 'Millonarios'
    if 'ARGENTINA' in n: return 'Argentina'
    if 'BRASIL' in n or 'BRAZIL' in n: return 'Brasil'
    if 'REAL MADRID' in n: return 'Real Madrid'
    if 'BARCELONA' in n: return 'Barcelona'
    if 'LIVERPOOL' in n: return 'Liverpool'
    if 'MANCHESTER' in n: return 'Manchester Utd'
    if 'INTER' in n: return 'Inter Milan'
    if 'JUNIOR' in n: return 'Junior'
    if 'AMERICA' in n: return 'América de Cali'
    if 'MEDELLIN' in n or 'MEDELLÍN' in n: return 'Medellín'
    if 'PORTUGAL' in n: return 'Portugal'
    if 'CHAPECOENSE' in n: return 'Chapecoense'
    if 'BAYERN' in n: return 'Bayern München'
    return 'Otro'

df_cam['calidad'] = df_cam['nombre'].apply(get_calidad)
df_cam['equipo'] = df_cam['nombre'].apply(get_equipo)

# Cálculos
total_fact = df['total'].sum()
total_cam = df_cam['total'].sum()
total_otros = total_fact - total_cam
unidades_cam = int(df_cam['cantidad'].sum())
ticket_prom = df_cam['total'].mean()

dias_top = df.groupby('fecha_dia')['total'].sum().sort_values(ascending=False)
dow = df.groupby('dia_semana')['total'].sum().rename(DIAS_ES)
dow_sorted = dow.sort_values(ascending=False)
horas = df.groupby('hora')['total'].sum().sort_values(ascending=False)
top_cam = df_cam.groupby('nombre').agg(unidades=('cantidad','sum'),ingreso=('total','sum')).sort_values('unidades',ascending=False)
cal = df_cam.groupby('calidad').agg(unidades=('cantidad','sum'),ingreso=('total','sum')).sort_values('ingreso',ascending=False)
eq = df_cam.groupby('equipo').agg(unidades=('cantidad','sum'),ingreso=('total','sum')).sort_values('ingreso',ascending=False)
top_general = df.groupby('nombre').agg(unidades=('cantidad','sum'),ingreso=('total','sum')).sort_values('ingreso',ascending=False)

# NEXY
df_n = pd.read_excel(os.path.join(BASE, '..', 'data', 'clientes',
    'HISTORIAL DE COMPRAS CLIENTES NEXY.xlsx'), header=6)
df_n.columns = ['cliente','nit','orden','fecha','productos']
df_n = df_n.dropna(subset=['cliente'])
df_n['fecha'] = pd.to_datetime(df_n['fecha'], errors='coerce')
df_n['mes'] = df_n['fecha'].dt.to_period('M')
ordenes_mes = df_n.groupby('mes').size().sort_index()
clientes_vip = df_n[~df_n['cliente'].str.upper().isin(['SIN TERCERO','CONSUMIDOR FINAL','ADDI','SISTECREDITO SAS','LLEVALO YA'])]['cliente'].value_counts().head(10)

# ══ Colores ══════════════════════════════════════════════════
VERDE  = RGBColor(0x00,0x6A,0x4E)
DORADO = RGBColor(0xD4,0xAF,0x37)
OSCURO = RGBColor(0x0D,0x1B,0x2A)
GRIS   = RGBColor(0x55,0x55,0x55)
ROJO   = RGBColor(0xC0,0x39,0x2B)
BLANCO = RGBColor(0xFF,0xFF,0xFF)
AZUL   = RGBColor(0x1A,0x5C,0x8A)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.5)

def shade(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex6)
    tcPr.append(shd)

def divider(color='006A4E'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'6')
    bt.set(qn('w:space'),'1'); bt.set(qn('w:color'),color)
    pBdr.append(bt); pPr.append(pBdr)

def h1(text, color=VERDE, size=18, space_b=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color
    return p

def h2(text, color=OSCURO, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color

def body(text, size=10.5, color=GRIS, bold=False, space_a=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_a)
    r = p.add_run(text); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold

def bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size=Pt(size); r.font.color.rgb=GRIS

def table(headers, rows, hdr_color='006A4E', col_w=None, alt=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; shade(c, hdr_color)
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=BLANCO
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri,row in enumerate(rows):
        tr = t.add_row()
        bg = 'F2F9F6' if (ri%2==0 and alt) else 'FFFFFF'
        for ci,val in enumerate(row):
            c = tr.cells[ci]; shade(c,bg)
            p = c.paragraphs[0]; r = p.add_run(str(val))
            r.font.size=Pt(9.5); r.font.color.rgb=OSCURO
    if col_w:
        for i,w in enumerate(col_w):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def kpi_box(label, value, sublabel='', color=VERDE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{label}: '); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=GRIS
    r2 = p.add_run(value); r2.bold=True; r2.font.size=Pt(13); r2.font.color.rgb=color
    if sublabel:
        r3 = p.add_run(f'  {sublabel}'); r3.font.size=Pt(9); r3.font.color.rgb=RGBColor(0x88,0x88,0x88)

# ══════════════════ PORTADA ══════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run('B370 LÍNEA DEPORTIVA'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=VERDE

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('INFORME DE VENTAS Y BASE DE CLIENTES')
r2.font.size=Pt(15); r2.font.color.rgb=OSCURO; r2.bold=True

p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Abril 4 – Mayo 6, 2026  |  Datos reales Quenti + NEXY')
r3.font.size=Pt(10); r3.font.color.rgb=GRIS; r3.italic=True

doc.add_paragraph()
divider('006A4E')

# 4 KPI grandes en portada
doc.add_paragraph()
kpi_box('Facturación total', f'$67,048,298 COP', '(32 días)', VERDE)
kpi_box('Ingresos por camisetas', f'$18,646,500 COP', f'({int(unidades_cam)} unidades vendidas)', VERDE)
kpi_box('Ticket promedio camiseta', f'${ticket_prom:,.0f} COP', '', AZUL)
kpi_box('Clientes únicos histórico', '2,101', '(Base NEXY mayo 2025 - marzo 2026)', AZUL)
kpi_box('Órdenes históricas', '5,399', '(NEXY mayo 2025 - marzo 2026)', AZUL)

doc.add_page_break()

# ══════════════════ 1. RESUMEN EJECUTIVO ═════════════════════════
h1('1. Resumen Ejecutivo', size=20)
body('Período analizado: 4 de abril al 6 de mayo de 2026 (datos operativos Quenti). '
     'Base de clientes histórica: mayo 2025 a marzo 2026 (sistema NEXY).', size=10.5)
divider()

table(
    ['Indicador', 'Valor', 'Notas'],
    [
        ['Facturación bruta total', '$67,048,298 COP', 'Todos los productos (camisetas + calzado + accesorios)'],
        ['Facturación camisetas', '$18,646,500 COP', '27.8% del total facturado'],
        ['Otros productos (guayos, balones, medias...)', '$48,401,798 COP', '72.2% del total'],
        ['Unidades de camisetas vendidas', f'{unidades_cam} und', 'Promedio ~33 camisetas/semana'],
        ['Precio promedio camiseta vendida', f'${ticket_prom:,.0f} COP', 'Refleja mix de calidades'],
        ['Período analizado (Quenti)', '32 días', '4 abril – 6 mayo 2026'],
        ['Clientes únicos (base histórica)', '2,101', 'Sistema NEXY, mayo 2025 – marzo 2026'],
        ['Órdenes históricas totales', '5,399', 'Sistema NEXY'],
    ],
    col_w=[7, 5, 6]
)

body('HALLAZGO CLAVE: La tienda física genera la mayor parte del ingreso en categorías de calzado deportivo '
     '(guayos, futsala) y accesorios. Las camisetas son el 28% de la facturación — pero son el diferenciador '
     'digital y el motor de la marca B370 online.', bold=True, color=OSCURO)

doc.add_page_break()

# ══════════════════ 2. DÍAS Y HORAS DE MÁS VENTAS ════════════════
h1('2. Días y Horas de Mayor Venta', size=20)
divider()

h2('Top 15 días con mayor facturación (abril–mayo 2026)')
rows_dias = []
for i, (fecha, val) in enumerate(dias_top.head(15).items()):
    bar = '█' * int(val / 500000)
    rows_dias.append([str(i+1), str(fecha.strftime('%d %b %Y')), f'${val:,.0f}', bar[:12]])
table(['#','Fecha','Facturación COP','Volumen relativo'], rows_dias,
      col_w=[1.2, 4, 5, 7.5])

body('El 2 de mayo fue el día más fuerte ($4,834,900). Los sábados dominan consistentemente la semana.', size=10)

h2('Facturación por día de la semana')
rows_dow = []
for dia, val in dow_sorted.items():
    pct = val / dow_sorted.sum() * 100
    bar = '█' * int(pct / 3)
    rows_dow.append([dia, f'${val:,.0f}', f'{pct:.1f}%', bar])
table(['Día','Facturación COP','% del total','Volumen'], rows_dow, col_w=[4,5,3,6])

body('SÁBADO ES EL REY: representa más del doble que el domingo. '
     'El lunes también sorprende — posiblemente compras online del fin de semana que se cierran el lunes.', bold=True, color=VERDE)

h2('Horas pico de venta')
rows_h = []
for hora, val in horas.head(10).items():
    if hora == 0: continue
    bar = '█' * int(val / 700000)
    rows_h.append([f'{hora:02d}:00 – {hora+1:02d}:00', f'${val:,.0f}', bar[:14]])
table(['Hora','Facturación COP','Volumen'], rows_h, col_w=[5,5,8])

body('Las ventas se concentran en 3 ventanas: mañana (11am), tarde (3-4pm) y tarde-noche (6pm). '
     'Estos son los mejores momentos para publicar en redes y enviar emails.', bold=True, color=VERDE)

doc.add_page_break()

# ══════════════════ 3. CAMISETAS MÁS VENDIDAS ══════════════════════
h1('3. Camisetas Más Vendidas', size=20)
divider()

h2('Top 20 camisetas por unidades vendidas')
rows_c = []
for i, (nombre, row) in enumerate(top_cam.head(20).iterrows()):
    # Limpiar nombre para display
    n_short = nombre.replace('CAMISETA ','').replace('DE HOMBRE','Hombre').replace('DE DAMA','Dama')
    n_short = n_short[:65] + ('...' if len(n_short) > 65 else '')
    rows_c.append([str(i+1), n_short, str(int(row['unidades'])), f'${row["ingreso"]:,.0f}'])
table(['#','Producto','Und.','Ingreso COP'], rows_c, col_w=[1.2, 10.5, 2, 4])

h2('Ventas por equipo / selección')
rows_eq = []
for equipo, row in eq.iterrows():
    pct = row['ingreso'] / cal['ingreso'].sum() * 100
    rows_eq.append([equipo, str(int(row['unidades'])), f'${row["ingreso"]:,.0f}', f'{pct:.1f}%'])
table(['Equipo/Selección','Unidades','Ingreso COP','% del total camisetas'],
      rows_eq, col_w=[5.5, 3, 5, 4])

body('Colombia es el producto estrella: 94 unidades y $7.7M en un mes. '
     'Con el Mundial 2026 en junio, esta línea tiene todo el potencial de triplicar volumen.', bold=True, color=VERDE)

h2('Ventas por calidad')
rows_cal = []
for calidad, row in cal.iterrows():
    pct_u = row['unidades'] / cal['unidades'].sum() * 100
    pct_i = row['ingreso'] / cal['ingreso'].sum() * 100
    rows_cal.append([calidad, str(int(row['unidades'])), f'{pct_u:.1f}%',
                     f'${row["ingreso"]:,.0f}', f'{pct_i:.1f}%'])
table(['Calidad','Unidades','% unidades','Ingreso COP','% ingresos'],
      rows_cal, col_w=[4, 3, 3, 5, 3])

body('Tipo Original domina en ingresos. Retro tiene muy buen volumen — es el producto con mejor '
     'relación precio/rotación. El 1.1 aún tiene poco volumen — hay oportunidad de posicionarlo mejor.', bold=True, color=OSCURO)

doc.add_page_break()

# ══════════════════ 4. TOP PRODUCTOS GENERALES ════════════════════
h1('4. Top Productos por Ingreso (Todo el catálogo)', size=20)
divider()
body('Incluye camisetas, guayos, futsala, balones, medias y accesorios. '
     'Permite ver el peso real de cada categoría en la facturación.', size=10)

rows_g = []
for i, (nombre, row) in enumerate(top_general.head(30).iterrows()):
    if 'ALBUM' in nombre.upper() or 'SOBRE' in nombre.upper(): continue
    n_short = nombre[:60] + ('...' if len(nombre) > 60 else '')
    rows_g.append([str(i+1), n_short, str(int(row['unidades'])), f'${row["ingreso"]:,.0f}'])
    if len(rows_g) >= 25: break
table(['#','Producto','Und.','Ingreso COP'], rows_g, col_w=[1.2, 11, 2, 4])

doc.add_page_break()

# ══════════════════ 5. BASE DE CLIENTES ══════════════════════════
h1('5. Base de Clientes (NEXY — Historial Completo)', size=20)
body('Período: mayo 2025 – marzo 2026. Total: 5,399 órdenes | 2,101 clientes únicos.', size=10.5)
divider()

h2('Crecimiento mensual de órdenes')
rows_mes = []
for mes, cnt in ordenes_mes.items():
    bar = '█' * int(cnt / 50)
    rows_mes.append([str(mes), str(cnt), bar[:18]])
table(['Mes','Órdenes','Volumen'], rows_mes, col_w=[4, 3, 11])

body('Diciembre 2025 fue el mes pico (1,143 órdenes). Hay estacionalidad clara: '
     'el Q4 (oct-dic) concentra la mayor actividad — navidad, año nuevo, temporada de ligas.', bold=True, color=VERDE)

h2('Top 10 clientes frecuentes (excluyendo "sin tercero" y pasarelas de pago)')
rows_vip = []
for cliente, cnt in clientes_vip.items():
    rows_vip.append([cliente, str(cnt)])
table(['Cliente','Órdenes'], rows_vip, col_w=[12, 3])

body('Estos son los clientes con mayor lealtad — candidatos directos a programa VIP, '
     'acceso anticipado a lanzamientos y bonos de fidelización.', bold=True, color=VERDE)

h2('Segmentación de la base de clientes')
table(
    ['Segmento','Criterio','Tamaño estimado','Estrategia'],
    [
        ['Clientes VIP','+5 compras históricas','~50-80 clientes','Programa VIP, bono ORO/PLATA, acceso anticipado'],
        ['Clientes recurrentes','2-4 compras','~300 clientes','Campaña de reactivación, cupón VUELVE10/VUELVE20'],
        ['Compradores únicos','1 compra','~1,700 clientes','Secuencia de email post-compra → segunda compra'],
        ['Leads WhatsApp','Consulta sin compra','En proceso','Follow-up activo, resolver objeción de talla/precio'],
    ],
    col_w=[4, 4.5, 4, 5.5]
)

doc.add_page_break()

# ══════════════════ 6. RECOMENDACIONES ══════════════════════════
h1('6. Recomendaciones Estratégicas', size=20, color=ROJO)
divider('C0392B')

h2('Inmediatas (esta semana)')
bullet('Activar campaña Colombia Mundial: el producto más vendido en un período donde el Mundial se acerca. '
       'Cada semana que pase sin pauta activa es dinero que se va.')
bullet('Programar publicaciones en horarios pico: 11am, 3pm y 6pm — especialmente los sábados.')
bullet('Activar secuencia de carrito abandonado en Omnisend: los 1,700 clientes de compra única son '
       'el segmento más fácil de convertir a segunda compra.')

h2('Esta quincena')
bullet('Lanzar programa VIP: los ~80 clientes con 5+ compras merecen bono ORO/PLATA y acceso anticipado. '
       'Cuesta poco y genera lealtad muy alta.')
bullet('Empujar camiseta 1.1: solo 5 unidades vendidas en el período — está completamente sub-comunicada. '
       'Es el producto de mayor margen.')
bullet('Activar Colombia Visitante 2026: tercera más vendida en unidades pero casi sin pauta propia.')

h2('Este mes')
bullet('Preparar inventario de Colombia para el Mundial (11 junio): basado en los datos, Colombia '
       'representa el 44% del volumen de camisetas. El pico del Mundial puede duplicar esa demanda.')
bullet('Activar campaña de reviews Google: el mayor riesgo identificado es la falta de reseñas. '
       'Los 2,101 clientes históricos son el ejército para lograrlo.')
bullet('Analizar qué compran junto con las camisetas: guayos y medias son la categoría más grande. '
       'Un bundle camiseta + medias puede aumentar el ticket promedio fácilmente.')

divider()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fin.paragraph_format.space_before = Pt(30)
r_fin = p_fin.add_run('B370 Línea Deportiva — La Ceja, Antioquia — b370sports.com')
r_fin.font.size=Pt(9); r_fin.font.color.rgb=GRIS; r_fin.italic=True

p_fin2 = doc.add_paragraph()
p_fin2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin2 = p_fin2.add_run('Datos: Quenti (abril–mayo 2026) + NEXY (mayo 2025–marzo 2026)  |  Confidencial')
r_fin2.font.size=Pt(8); r_fin2.font.color.rgb=RGBColor(0xAA,0xAA,0xAA); r_fin2.italic=True

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f'Generado: {os.path.abspath(OUTPUT)}')
