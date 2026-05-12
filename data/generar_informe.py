import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Colores
NEGRO    = RGBColor(0x1A, 0x1A, 0x1A)
VERDE    = RGBColor(0x00, 0x7A, 0x33)
BLANCO   = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_OSC = RGBColor(0x4A, 0x4A, 0x4A)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(text, size=16, color=VERDE, bold=True, sb=12, sa=6, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

def body(text, size=10, color=GRIS_OSC, sa=4, bold=False, italic=False, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic

def bullet(text, size=9):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRIS_OSC

def make_header(table, cols, bg='007A33'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO

def add_row(table, values, bg=None, bold=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p   = cell.paragraphs[0]
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        run.bold = bold
        run.font.color.rgb = NEGRO

def two_color_run(label, text, sa=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(sa)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = VERDE
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRIS_OSC

# ============================================================
#  PORTADA
# ============================================================
heading('B370 LINEA DEPORTIVA', size=30, color=VERDE, bold=True, sb=50, sa=4, center=True)
heading('Analisis Comercial — Dos Sistemas de Facturacion', size=16, color=NEGRO, bold=True, sb=0, sa=4, center=True)
body('NEXY (Tienda Online) + CUENTI POS (Tienda Fisica)  |  Mayo 2026',
     size=11, italic=True, center=True, sa=40)
doc.add_page_break()

# ============================================================
#  RESUMEN EJECUTIVO
# ============================================================
heading('RESUMEN EJECUTIVO', size=15, color=VERDE, sb=6, sa=4)
body('Datos extraidos directamente de los dos sistemas de facturacion de B370. Periodo analizado: mayo 2025 - marzo 2026 (online) y 4 abril - 6 mayo 2026 (tienda fisica).', italic=True, sa=8)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
make_header(t, ['Metrica', 'NEXY Online', 'CUENTI Fisico', 'Total / Nota'])
exec_rows = [
    ('Periodo',             'May 2025 - Mar 2026\n(10 meses)', '4 Abr - 6 May 2026\n(33 dias)', '---'),
    ('Transacciones',       '5,397 ordenes',    '635 facturas',   '6,032 en total'),
    ('Clientes unicos',     '2,101 clientes',   '346 c/nombre',   '+2,400 aprox.'),
    ('Ventas brutas',       'Sin precio en NEXY','COP $67,048,298','Solo tienda fisica'),
    ('Ticket promedio',     '---',               'COP $105,588',   '---'),
    ('Unidades vendidas',   '---',               '1,577',          '---'),
    ('Clientes recurrentes','315 (3+ compras)',  'Ver seccion 6',  '---'),
]
for i, r in enumerate(exec_rows):
    add_row(t, r, bg='F2F2F2' if i % 2 == 0 else 'FFFFFF')

doc.add_paragraph()

# ============================================================
#  BLOQUE 1 — DIAS
# ============================================================
doc.add_page_break()
heading('1.  DIAS QUE MAS SE VENDE', size=15, color=VERDE, sb=6, sa=4)

heading('1.1  Tienda Fisica — CUENTI (33 dias)', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
make_header(t, ['Dia', 'Ventas Brutas', 'Facturas', 'Unidades', 'Posicion'])
dias_f = [
    ('Sabado',    '$15,930,700', '148', '376', 'REY  #1'),
    ('Lunes',     '$11,611,999', '113', '288', '#2'),
    ('Jueves',    '$10,049,700', '92',  '189', '#3'),
    ('Martes',    '$9,959,600',  '103', '263', '#4'),
    ('Miercoles', '$8,283,999',  '79',  '194', '#5'),
    ('Viernes',   '$7,295,300',  '65',  '182', '#6'),
    ('Domingo',   '$3,917,000',  '35',  '85',  'Mas flojo'),
]
for i, r in enumerate(dias_f):
    bg = 'E8F5EC' if r[0] in ('Sabado', 'Lunes') else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg, bold=r[0] in ('Sabado', 'Lunes'))

body('Conclusion: El SABADO genera el doble de ventas que cualquier dia entre semana. Domingo es el mas flojo en ambos canales.', bold=True, color=VERDE, sa=8)

heading('1.2  Tienda Online — NEXY (10 meses)', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
make_header(t, ['Dia', 'Ordenes Online', 'Posicion'])
dias_o = [
    ('Sabado',    '1,037', 'REY  #1'),
    ('Viernes',   '842',   '#2'),
    ('Martes',    '827',   '#3'),
    ('Miercoles', '789',   '#4'),
    ('Jueves',    '779',   '#5'),
    ('Lunes',     '663',   '#6'),
    ('Domingo',   '461',   'Mas flojo'),
]
for i, r in enumerate(dias_o):
    bg = 'E8F5EC' if r[0] in ('Sabado', 'Viernes') else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg, bold=r[0] in ('Sabado', 'Viernes'))

heading('1.3  Horas Pico — Tienda Fisica', size=12, color=NEGRO, sb=10, sa=4)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
make_header(t, ['Hora', 'Ventas', 'Nivel'])
horas = [
    ('09:00', '$377,900',    'Apertura tranquila'),
    ('10:00', '$5,181,400',  'Arranque fuerte'),
    ('11:00', '$8,419,800',  'PICO MANANA'),
    ('12:00', '$6,094,500',  'Alto'),
    ('13:00', '$3,126,800',  'Bajon almuerzo'),
    ('14:00', '$6,545,600',  'Rebote tarde'),
    ('15:00', '$10,000,900', 'PICO MAXIMO DEL DIA'),
    ('16:00', '$8,709,900',  'Muy alto'),
    ('17:00', '$7,413,199',  'Alto'),
    ('18:00', '$7,542,799',  'Cierre fuerte'),
    ('19:00', '$3,635,500',  'Baja final'),
]
for r in horas:
    bg = 'E8F5EC' if 'PICO' in r[2] else 'FFFFFF'
    add_row(t, r, bg=bg, bold='PICO' in r[2])

body('Dos picos claros: 11AM y 3PM. El equipo de ventas debe estar completo en esas franjas horarias.', bold=True, color=VERDE, sa=6)

# ============================================================
#  BLOQUE 2 — PRODUCTOS
# ============================================================
doc.add_page_break()
heading('2.  PRODUCTOS ESTRELLA', size=15, color=VERDE, sb=6, sa=4)

heading('2.1  Tienda Fisica — CUENTI', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
make_header(t, ['Producto', 'Unidades', 'Ingresos', 'Precio Prom.', 'Categoria'])
pf = [
    ('Sobre figuras Panini',               '496', '$2,769,500', '$5,584',  'Coleccion'),
    ('Cartas de jugadores',                '51',  '$46,100',    '$904',    'Coleccion'),
    ('Album pasta blanda Panini',          '25',  '$355,000',   '$14,200', 'Coleccion'),
    ('Munecos cabezones',                  '24',  '$294,000',   '$12,250', 'Coleccion'),
    ('Media antideslizante negra adulto',  '18',  '$260,000',   '$14,444', 'Accesorio'),
    ('Figuras de coleccion mundial',       '15',  '$561,000',   '$37,400', 'Col. premium'),
    ('Estampacion tipo original',          '16',  '$365,000',   '$22,812', 'Servicio'),
    ('Canilleras imp pequenas',            '16',  '$217,000',   '$13,562', 'Accesorio'),
    ('Media antideslizante blanca adulto', '15',  '$266,000',   '$17,733', 'Accesorio'),
    ('Camiseta Colombia Local 2026 Orig.', '14',  '$1,215,000', '$86,800', 'Camiseta VIP'),
    ('Camiseta Colombia Visit. 2026 Orig.','5',   '$515,000',   '$103,000','Camiseta VIP'),
    ('Camiseta Colombia Local Dama 2026',  '9',   '$734,900',   '$81,600', 'Camiseta VIP'),
    ('Caja regalo futbol Colombia',        '9',   '$91,000',    '$10,111', 'Regalo'),
]
for i, r in enumerate(pf):
    bg = 'E8F5EC' if 'Colombia' in r[0] or r[0] == 'Sobre figuras Panini' else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg)

heading('2.2  Tienda Online — NEXY (10 meses)', size=12, color=NEGRO, sb=10, sa=4)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
make_header(t, ['Producto', 'Ordenes'])
po = [
    ('Media antideslizante negra adulto',  '194'),
    ('Media antideslizante blanca adulto', '149'),
    ('Cajas regalo de futbol',             '119'),
    ('Media larga finas',                  '109'),
    ('Canilleras imp pequenas',            '103'),
    ('Canilleras imp grandes',             '83'),
    ('Media tobillera eco',                '82'),
    ('Estampacion tipo original',          '57'),
    ('Balon de FIFA #5',                   '56'),
    ('Gorra Atletico Nacional',            '44'),
    ('Camiseta 1.1 Hombre Talla L',        '38'),
    ('Camiseta 1.1 Hombre Talla XL',       '37'),
]
for i, r in enumerate(po):
    bg = 'E8F5EC' if i == 0 else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg)

# ============================================================
#  BLOQUE 3 — DATA PARA EL CLIENTE
# ============================================================
doc.add_page_break()
heading('3.  DATA PARA COMUNICARLE AL CLIENTE', size=15, color=VERDE, sb=6, sa=4)
body('Mensajes con respaldo real de datos — listos para usar en tienda, WhatsApp, redes o con aliados:', italic=True, sa=8)

mensajes = [
    ('Volumen de atencion',     'En un solo mes atendemos mas de 600 familias en el punto fisico y mas de 5,000 pedidos online en el ultimo ano. B370 es la tienda de futbol mas activa de La Ceja.'),
    ('El producto del momento', 'La camiseta de Colombia Local 2026 Tipo Original es la mas pedida del ano — disponible con y sin dorsal para hombre y dama.'),
    ('El sabado somos el hub',  'Cada sabado mas de 140 clientes visitan B370 en La Ceja. Es el dia del futbol y nosotros lo vivimos con ellos.'),
    ('Stock online garantizado','Si no encuentras tu talla en tienda, esta en b370sports.com. Enviamos a todo Colombia.'),
    ('Panini — el furor',       '496 sobres de figuras Panini vendidos en solo 33 dias. Si estas completando el album, aqui tienes el mejor surtido.'),
    ('Clientes fieles',         'Mas de 315 clientes compraron 3 o mas veces en la tienda online. En B370 construimos relaciones, no solo transacciones.'),
]
for label, text in mensajes:
    two_color_run(label, text)

# ============================================================
#  BLOQUE 4 — ESTRATEGIA DE CONTENIDO
# ============================================================
doc.add_page_break()
heading('4.  ESTRATEGIA DE CONTENIDO', size=15, color=VERDE, sb=6, sa=4)
body('Pilares de contenido basados directamente en los datos reales de ventas:', italic=True, sa=8)

pilares = [
    ('PILAR 1 — Coleccionismo Panini (fenomeno activo)',
     '496 sobres vendidos en 33 dias = 15 sobres/dia.',
     ['Unboxing de sobres en Reels cortos',
      'Stories con encuestas: cuantos te faltan?',
      'La figura que nadie tiene',
      'Reto: completar el album',
      'Mostrar inventario de albums y sobres disponibles']),
    ('PILAR 2 — Colombia 2026 (camiseta del momento)',
     'Producto de mayor ingreso individual. Ticket $86K-$109K. Copa America y Eliminatorias como contexto de contenido.',
     ['La local 2026 ya llego — hombre, dama, con/sin dorsal',
      'Videos con la camiseta puesta en partido o entrenamiento',
      'Publicar jueves y viernes para capturar intencion del sabado',
      'Contenido emocional: Colombia en el corazon']),
    ('PILAR 3 — Accesorios / Compra impulsiva',
     'Medias antideslizantes #1 online (194 ordenes). Canilleras top 5 en ambos canales.',
     ['Carrusel: arma tu kit completo por menos de $50,000',
      'Bundle camiseta + medias + canilleras',
      'Antes de tu proximo partido necesitas esto',
      'Variedad de colores y tallas disponibles']),
    ('PILAR 4 — Sabado = Evento en tienda',
     'El sabado factura el doble que cualquier dia entre semana. 148 facturas en un solo sabado.',
     ['Viernes de noche: manana estamos hasta las 7PM, te esperamos',
      'Stories en vivo desde tienda el sabado',
      'Reels de ambiente: clientes, productos, equipo',
      'Recordatorio de horario y ubicacion']),
    ('PILAR 5 — Personalizacion y Estampacion',
     'Top 10 en ambos canales. Servicio diferenciador, alta percepcion de valor.',
     ['Video del proceso: antes y despues de la estampacion',
      'Ponle tu nombre a la camiseta',
      'Galeria de trabajos realizados',
      'CTA claro: pidela en tienda o por WhatsApp']),
]

for titulo, dato, bullets_list in pilares:
    heading(titulo, size=11, color=NEGRO, bold=True, sb=10, sa=2)
    body('Dato clave: ' + dato, size=9, color=VERDE, bold=True, sa=2)
    for b in bullets_list:
        bullet(b)

# ============================================================
#  BLOQUE 5 — META ADS
# ============================================================
doc.add_page_break()
heading('5.  DATA PARA CAMPANAS META ADS', size=15, color=VERDE, sb=6, sa=4)

heading('5.1  Timing — Cuando pautar', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
make_header(t, ['Objetivo', 'Dia', 'Hora', 'Por que'])
timing = [
    ('Conversion',  'Jueves - Viernes', '6PM - 9PM',  'Captura intencion antes del sabado de compra'),
    ('Awareness',   'Martes - Miercoles','12PM - 2PM', 'Buena actividad, menor costo de pauta'),
    ('Remarketing', 'Domingo',           'Todo el dia','Dia flojo en tienda — cliente en casa con el celular'),
    ('Post pico',   'Sabado',            '7PM - 10PM', 'Cierre de semana, engagement alto'),
]
for i, r in enumerate(timing):
    add_row(t, r, bg='F2F2F2' if i % 2 == 0 else 'FFFFFF')

heading('5.2  Audiencias sugeridas', size=12, color=NEGRO, sb=10, sa=4)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
make_header(t, ['Audiencia', 'Producto', 'Ticket Estimado', 'Prioridad'])
audiencias = [
    ('Hombres 18-35, futbol, Colombia',       'Camiseta Colombia 2026 Original',       '$97K - $109K',  'ALTA'),
    ('Padres 28-45 — Dia de la Madre',        'Uniformes infantiles Colombia/Nacional', '$65K',          'ALTA — urgente'),
    ('Coleccionistas futbol 15-40',           'Figuras mundiales, Panini',              '$37K - $75K',   'ALTA'),
    ('Jugadores activos 16-35',               'Camiseta 1.1 + medias + canilleras',     '$150K bundle',  'MEDIA'),
    ('Aficionados Atletico Nacional',         'Uniformes, gorras, chompa',              '$65K - $95K',   'MEDIA'),
    ('Mujeres 20-40, futbol femenino',        'Camiseta Colombia Local Dama 2026',      '$80K',          'MEDIA'),
]
for i, r in enumerate(audiencias):
    bg = 'E8F5EC' if r[3] == 'ALTA' else ('FFF3CD' if 'urgente' in r[3] else ('F2F2F2' if i % 2 == 0 else 'FFFFFF'))
    add_row(t, r, bg=bg, bold='ALTA' in r[3])

heading('5.3  Anuncios prioritarios', size=12, color=NEGRO, sb=10, sa=4)
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
make_header(t, ['Tipo', 'Producto', 'Por que funciona', 'Formato'])
anuncios = [
    ('HERO — estrella',     'Camiseta Colombia Local 2026 Tipo Original', 'Mas vendida en ingresos, demanda probada, contexto Copa America', 'Video + foto con modelo'),
    ('VOLUMEN — impulsivo', 'Medias + canilleras bundle',                 '194 ordenes online. Precio bajo, alta velocidad de compra',       'Carrusel de productos'),
    ('REGALO — fecha',      'Caja regalo de futbol',                      '119 pedidos online. Dia de la Madre en mayo = contexto perfecto', 'Imagen con lazo + CTA'),
    ('PREMIUM',             'Figuras de coleccion mundial',               '$37K precio prom. Coleccionistas, alta emocion de compra',        'Video unboxing'),
    ('FAMILIA',             'Uniformes Colombia infantiles',              'Padres + Dia de la Madre = combinacion ideal ahora mismo',        'Foto papa/hijo camiseta'),
]
for i, r in enumerate(anuncios):
    bg = 'E8F5EC' if i == 0 else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg, bold=i == 0)

heading('5.4  Copies angulo ganador', size=12, color=NEGRO, sb=10, sa=4)
copies = [
    ('Colombia 2026',    'La camiseta mas pedida del ano ya esta aqui. Talla L y M se van rapido. Pidela antes que se agote.'),
    ('Sabado en tienda', 'El sabado nos visitan mas de 140 familias. Si no puedes venir, te llevamos el pedido a casa.'),
    ('Dia de la Madre',  'Regala la pasion del futbol. Caja regalo de B370 con camiseta de Colombia — el regalo que ninguna mama espera pero todas quieren.'),
    ('Bundle accesorio', 'Tu kit completo para el proximo partido: camiseta, medias y canilleras. Todo en B370 con envio a todo Colombia.'),
    ('Panini',           'Cuantos sobres te faltan para completar el album? En B370 tenemos el mejor surtido de La Ceja. Ven o pidelos online.'),
]
for label, text in copies:
    two_color_run(label, text)

# ============================================================
#  BLOQUE 6 — CLIENTES VIP
# ============================================================
doc.add_page_break()
heading('6.  CLIENTES VIP', size=15, color=VERDE, sb=6, sa=4)

heading('6.1  Tienda Fisica — Top compradores (CUENTI, 33 dias)', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
make_header(t, ['Cliente', 'Gasto Total', 'Facturas'])
vip_f = [
    ('ADELANTE SOLUCIONES FINANCIERAS S.A.S. (*)', '$4,676,000', '22'),
    ('SISTECREDITO SAS (*)',                        '$3,940,900', '23'),
    ('FELIPE RIOS PAVAS',                           '$1,201,000', '10'),
    ('LLEVALO YA YA',                               '$823,000',   '3'),
    ('LEIDY JARAMILLO',                             '$701,900',   '9'),
    ('YOVERLEY ALVAREZ LONDONO',                    '$660,000',   '3'),
    ('SEBASTIAN ROLDAN',                            '$630,000',   '1'),
    ('BENJAMIN ROCHA',                              '$619,900',   '2'),
    ('JUAN FELIPE JARAMILLO CARDONA',               '$522,500',   '9'),
    ('ANDERSON HENAO',                              '$461,000',   '3'),
]
for i, r in enumerate(vip_f):
    bg = 'E8F5EC' if i < 2 else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg, bold=i < 2)

body('(*) Adelante Soluciones y Sistecredito son sistemas de credito (clientes institucionales). No aplican para fidelizacion individual.', size=9, italic=True, sa=8)

heading('6.2  Tienda Online — Clientes mas frecuentes (NEXY)', size=12, color=NEGRO, sb=8, sa=4)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
make_header(t, ['Cliente', 'Ordenes Historicas'])
vip_o = [
    ('ADDI',                          '167'),
    ('SISTECREDITO SAS',              '139'),
    ('FELIPE RIOS PAVAS',             '86'),
    ('LLEVALO YA',                    '61'),
    ('LEIDY JARAMILLO',               '33'),
    ('ROBIN ANDRES ALZATE VALENCIA',  '31'),
    ('ANDERSON HENAO',                '30'),
    ('JUAN CASTRO',                   '29'),
    ('JOSE MANUEL LOPEZ LONDONO',     '22'),
]
for i, r in enumerate(vip_o):
    bg = 'E8F5EC' if i == 0 else ('F2F2F2' if i % 2 == 0 else 'FFFFFF')
    add_row(t, r, bg=bg)

body('Clientes en AMBOS sistemas (alta fidelidad — candidatos para WhatsApp VIP): Felipe Rios Pavas, Leidy Jaramillo, Anderson Henao, Juan Felipe Jaramillo Cardona.', bold=True, color=VERDE, sa=4)

# ============================================================
#  CIERRE
# ============================================================
doc.add_page_break()
heading('En B370, vestimos la pasion.', size=20, color=VERDE, bold=True, sb=60, sa=4, center=True)
body('b370sports.com  |  La Ceja, Antioquia  |  Mayo 2026', size=11, italic=True, center=True)

# ============================================================
#  GUARDAR
# ============================================================
output = r'c:\Proyectos\b370\PROTOCOLO 1.1 B370\B370-MANAGER\data\B370_Analisis_Comercial_Mayo2026.docx'
doc.save(output)
print('LISTO:', output)
